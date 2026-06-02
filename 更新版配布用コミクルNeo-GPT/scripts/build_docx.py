#!/usr/bin/env python3
"""
manuscript.md + manga_compiled.md + 画像 → final_book.docx
"""
import io
import re
import json
import sys
from pathlib import Path

# Windows cp932 対策: stdout/stderrをUTF-8で強制出力
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SLUG = "gengo-ryoku"
BASE = Path(__file__).parent.parent
OUT_DIR = BASE / "output" / SLUG

def add_page_break(doc):
    doc.add_page_break()

def set_heading_style(para, level, color_hex="#4A6CF7"):
    run = para.runs[0] if para.runs else para.add_run()
    run.bold = True
    r, g, b = int(color_hex[1:3],16), int(color_hex[3:5],16), int(color_hex[5:7],16)
    run.font.color.rgb = RGBColor(r, g, b)
    if level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(16)
    else:
        run.font.size = Pt(13)

def add_runs_with_bold(para, text):
    """
    **太字** を実際のbold runに変換してパラグラフに追加する。
    `**` の前後に記号・句読点があっても正しく処理する。
    """
    # `code` を除去してテキストのみに
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # **...** を分割して処理
    # 前後に記号がついていてもマッチするよう \*\* を単純に分割
    parts = re.split(r'\*\*', text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        # 奇数インデックス（**で囲まれた部分）をbold
        if idx % 2 == 1:
            run.bold = True


def is_table_line(line):
    """Markdown表の行かどうか判定"""
    return line.strip().startswith("|") and line.strip().endswith("|")

def is_separator_line(line):
    """表のセパレータ行（|---|---|）かどうか判定"""
    return bool(re.match(r'^\s*\|[\s\-|:]+\|\s*$', line))

def parse_table_cells(line):
    """表の1行をセルのリストに分割"""
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]

def insert_table(doc, table_lines):
    """Markdown表をWordテーブルに変換して挿入"""
    # セパレータ行を除外してデータ行だけ残す
    data_lines = [l for l in table_lines if not is_separator_line(l)]
    if not data_lines:
        return

    rows = [parse_table_cells(l) for l in data_lines]
    col_count = max(len(r) for r in rows)

    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.style = "Table Grid"

    for r_idx, row_cells in enumerate(rows):
        for c_idx in range(col_count):
            cell_text = row_cells[c_idx] if c_idx < len(row_cells) else ""
            cell = tbl.cell(r_idx, c_idx)
            para = cell.paragraphs[0]
            add_runs_with_bold(para, cell_text)
            # ヘッダー行は太字
            if r_idx == 0:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x4A, 0x6C, 0xF7)

    doc.add_paragraph()  # テーブル後に空行


def process_manuscript(doc, md_text, images_dir):
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Markdown表の検出・変換
        if is_table_line(line):
            table_lines = []
            while i < len(lines) and (is_table_line(lines[i]) or is_separator_line(lines[i])):
                table_lines.append(lines[i])
                i += 1
            insert_table(doc, table_lines)
            continue

        # ページブレーク
        if line.strip() == "\\newpage":
            add_page_break(doc)
            i += 1
            continue

        # 画像リンク ![alt](path)
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if img_match:
            alt, img_path = img_match.group(1), img_match.group(2)
            full_path = OUT_DIR / img_path
            if full_path.exists():
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(str(full_path), width=Inches(5.5))
                cap = doc.add_paragraph(alt)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].font.size = Pt(9)
                cap.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            i += 1
            continue

        # コメント（画像生成失敗など）をスキップ
        if line.strip().startswith("<!--"):
            i += 1
            continue

        # H1
        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip()
            if text:
                para = doc.add_heading(text, level=1)
                set_heading_style(para, 1)
            i += 1
            continue

        # H2
        if line.startswith("## "):
            text = line[3:].strip()
            if text:
                para = doc.add_heading(text, level=2)
                set_heading_style(para, 2)
            i += 1
            continue

        # H3
        if line.startswith("### "):
            text = line[4:].strip()
            if text:
                para = doc.add_heading(text, level=3)
                set_heading_style(para, 3)
            i += 1
            continue

        # 水平線
        if re.match(r'^---+$', line.strip()):
            i += 1
            continue

        # 箇条書き
        if line.strip().startswith("- "):
            text = line.strip()[2:]
            para = doc.add_paragraph(style='List Bullet')
            add_runs_with_bold(para, text)
            i += 1
            continue

        # 番号付きリスト
        num_match = re.match(r'^\d+\.\s+(.+)', line.strip())
        if num_match:
            text = num_match.group(1)
            para = doc.add_paragraph(style='List Number')
            add_runs_with_bold(para, text)
            i += 1
            continue

        # コードブロック
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                para = doc.add_paragraph("\n".join(code_lines))
                para.style = 'No Spacing'
                for run in para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 通常段落（太字を実際のbold runに変換）
        if line.strip():
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(14)
            para.paragraph_format.line_spacing = Pt(20)
            add_runs_with_bold(para, line.strip())

        i += 1


def insert_manga_pages(doc, page_numbers):
    """漫画画像を1画像=1ページフル表示で挿入する"""
    panels_dir = OUT_DIR / "panels"

    # A4ページ（余白0.5inch）の実効サイズ
    # 幅: 8.27 - 0.5*2 = 7.27inch  高さ: 11.69 - 0.5*2 = 10.69inch
    # 漫画アスペクト比 896:1200 = 0.747
    # width=7.27 → height=7.27/0.747=9.73inch ≤ 10.69 → OK
    MAX_WIDTH = Inches(7.27)
    MAX_HEIGHT = Inches(10.69)

    for n in page_numbers:
        add_page_break(doc)
        img_path = panels_dir / f"page_{n:02d}.png"
        if not img_path.exists():
            doc.add_paragraph(f"[漫画ページ {n} が見つかりません]")
            continue

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 段落の上下スペースをゼロに
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run()
        run.add_picture(str(img_path), width=MAX_WIDTH)


def split_manuscript_by_chapter(manuscript_text):
    """
    manuscript.md を章ごとに分割して返す。
    戻り値: [(section_name, text), ...]
    section_name は "intro", "ch1"〜"ch5", "outro" のいずれか
    """
    sections = []
    current_name = "intro"
    current_lines = []

    for line in manuscript_text.splitlines(keepends=True):
        # H1 の章見出しを検出（"# 第N章" または "# はじめに" / "# おわりに"）
        h1_match = re.match(r'^# (.+)', line)
        if h1_match:
            title = h1_match.group(1).strip()
            # 直前のセクションを確定
            if current_lines:
                sections.append((current_name, "".join(current_lines)))
            current_lines = [line]
            # セクション名を決定
            ch_match = re.search(r'第([1-5１-５])章', title)
            if ch_match:
                num = ch_match.group(1)
                # 全角→半角
                num = str("１２３４５".index(num) + 1) if num in "１２３４５" else num
                current_name = f"ch{num}"
            elif "おわりに" in title:
                current_name = "outro"
            else:
                current_name = "intro"
        else:
            current_lines.append(line)

    if current_lines:
        sections.append((current_name, "".join(current_lines)))

    return sections


# 章番号 → 漫画ページ番号のマッピング
CHAPTER_MANGA_PAGES = {
    "ch1": [1, 2, 3],
    "ch2": [4, 5, 6],
    "ch3": [7, 8, 9],
    "ch4": [10, 11, 12],
    "ch5": [13, 14, 15],
}


def main():
    meta = json.loads((OUT_DIR / "meta.json").read_text(encoding="utf-8"))
    manuscript_text = (OUT_DIR / "manuscript.md").read_text(encoding="utf-8")

    doc = Document()

    # ページ設定（A4）
    section = doc.sections[0]
    section.page_width = int(8.27 * 914400)
    section.page_height = int(11.69 * 914400)
    section.left_margin = int(0.5 * 914400)
    section.right_margin = int(0.5 * 914400)
    section.top_margin = int(0.5 * 914400)
    section.bottom_margin = int(0.5 * 914400)

    # 表紙
    doc.add_paragraph()
    title_para = doc.add_paragraph(meta["title"])
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0x4A, 0x6C, 0xF7)

    sub_para = doc.add_paragraph(meta["subtitle"])
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.runs[0]
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    target_para = doc.add_paragraph(f"対象読者: {meta['target_reader']}")
    target_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    target_para.runs[0].font.size = Pt(11)

    add_page_break(doc)

    # 原稿を章ごとに分割して処理
    print("原稿を章ごとに処理中...")
    sections = split_manuscript_by_chapter(manuscript_text)

    for section_name, section_text in sections:
        process_manuscript(doc, section_text, OUT_DIR / "images")

        # 第1章〜第5章の末尾に対応する漫画ページを挿入
        if section_name in CHAPTER_MANGA_PAGES:
            pages = CHAPTER_MANGA_PAGES[section_name]
            print(f"  {section_name} 漫画挿入: page_{pages[0]:02d}〜page_{pages[-1]:02d}")
            insert_manga_pages(doc, pages)
            add_page_break(doc)

    # 保存
    out_path = OUT_DIR / "final_book.docx"
    doc.save(str(out_path))
    size_mb = out_path.stat().st_size / 1024 / 1024
    print(f"\n完成: {out_path}")
    print(f"ファイルサイズ: {size_mb:.1f} MB")


if __name__ == "__main__":
    main()
